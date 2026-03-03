from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from core.document_creator.assemblers.base import BaseDocumentAssembler
from core.document_creator.state import DocumentCreatorConfig, SectionState


class DocxAssembler(BaseDocumentAssembler):
    """
    Assembles sections into a Word document using python-docx.

    Creates a professional report with:
    - Title page
    - Table of contents placeholder
    - Styled sections with headings, body text, bullets, and tables
    """

    async def assemble(
        self,
        title: str,
        subtitle: Optional[str],
        sections: List[SectionState],
        config: DocumentCreatorConfig,
        output_path: str,
    ) -> str:
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Title page
        self._add_title_page(doc, title, subtitle, config)

        # Sections
        for section_state in sections:
            if not section_state.versions:
                continue
            version = section_state.versions[section_state.selected_version_index]
            self._add_section(doc, section_state, version)

        doc.save(output_path)
        return output_path

    def _add_title_page(self, doc, title, subtitle, config):
        # Add some spacing
        for _ in range(6):
            doc.add_paragraph("")

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        # Subtitle
        if subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # Metadata
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Type: {config.document_type.value.replace('_', ' ').title()} | "
            f"Audience: {config.audience.value.title()}"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        doc.add_page_break()

    def _add_section(self, doc, section_state, version):
        heading_text = version.heading or section_state.spec.title
        level = min(section_state.spec.heading_level, 3)

        # Heading
        heading = doc.add_heading(heading_text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        # Key takeaway callout
        if version.key_takeaway:
            p = doc.add_paragraph()
            run = p.add_run(f"Key Takeaway: {version.key_takeaway}")
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xF9, 0x73, 0x16)
            p.paragraph_format.space_after = Pt(8)

        # Main content
        if version.content:
            for para_text in version.content.split("\n"):
                text = para_text.strip()
                if text:
                    doc.add_paragraph(text)

        # Bullet points
        if version.bullet_points:
            for bullet in version.bullet_points:
                doc.add_paragraph(bullet, style="List Bullet")

        # Table
        if version.table_data:
            self._add_table(doc, version.table_data)

        # Spacing after section
        doc.add_paragraph("")

    def _add_table(self, doc, table_data):
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        if not headers:
            return

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
            for p in header_cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True

        # Data rows
        for row_idx, row_data in enumerate(rows):
            cells = table.rows[row_idx + 1].cells
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < len(headers):
                    cells[col_idx].text = str(cell_value)
